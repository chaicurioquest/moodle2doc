import os
import xml.etree.ElementTree as ET
from docx import Document
from bs4 import BeautifulSoup

def parse_moodle_xml(xml_file):
    tree = ET.parse(xml_file)
    root = tree.getroot()
    questions = []

    for question in root.findall('question'):
        qtype = question.get('type')
        name = question.find('name').find('text').text
        questiontext_html = question.find('questiontext').find('text').text

        # Clean the HTML content
        questiontext = clean_html(questiontext_html)

        questions.append({
            'type': qtype,
            'name': name,
            'questiontext': questiontext
        })

    return questions

def clean_html(html_text):
    soup = BeautifulSoup(html_text, 'html.parser')
    return soup.get_text()

def create_word_document(questions, output_file, user_choices):
    doc = Document()
    doc.add_heading('Moodle Questions', 0)

    for q in questions:
        print(f"Processing question: {q['name']}")
        doc.add_heading(q['name'], level=1)
        doc.add_paragraph(f"Type: {q['type']}")
        doc.add_paragraph(f"Question: {q['questiontext']}")

    # Save the document at the specified path
    print(f"Saving document to: {output_file}")
    doc.save(output_file)
    print(f"Document saved successfully to: {output_file}")

def generate_document(xml_file, output_file, user_choices):
    try:
        print(f"Parsing XML file: {xml_file}")
        questions = parse_moodle_xml(xml_file)
        print(f"Parsed questions: {questions}")

        print(f"Creating Word document at: {output_file}")
        create_word_document(questions, output_file, user_choices)

        print(f"Document generation completed.")
    except Exception as e:
        print(f"Error generating document: {e}")
